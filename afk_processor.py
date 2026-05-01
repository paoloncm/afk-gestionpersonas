import os
import re
import json
import base64
from typing import Optional, List
import fitz
from pydantic import BaseModel, Field
from pypdf import PdfReader
import docx
from openai import OpenAI
from supabase import create_client, Client
from dotenv import load_dotenv
from datetime import datetime

load_dotenv()

# --- Schemas ---

class CandidateCV(BaseModel):
    nombre_completo: str = Field(..., description="Nombre completo del candidato")
    rut: Optional[str] = Field(None, description="RUT chileno si está presente, formato: 12.345.678-9")
    fecha_nacimiento: Optional[str] = Field(None, description="Fecha de nacimiento si está presente (formato YYYY-MM-DD)")
    profesion: Optional[str] = Field(None, description="Profesión principal o título")
    correo: Optional[str] = Field(None, description="Correo electrónico")
    telefono: Optional[str] = Field(None, description="Teléfono de contacto")
    direccion: Optional[str] = Field(None, description="Dirección física o ciudad de residencia")

    cargo: Optional[str] = Field(None, description="Cargo actual o más reciente")
    ultima_exp_laboral_empresa: Optional[str] = Field(None, description="Empresa más reciente")
    periodo: Optional[str] = Field(None, description="Período del empleo más reciente (ej: '2019 - Presente')")
    software_que_domina: Optional[str] = Field(None, description="Lista de software, ERPs (SAP, Maximo), herramientas y tecnologías que el candidato domina. Incluir todo.")
    experiencia: Optional[str] = Field(None, description="Resumen de la experiencia más reciente")

    cargo_a_desempenar: Optional[str] = Field(None, description="Cargo al que postula o último cargo objetivo")
    experiencia_total: int = Field(0, description="Total de años de experiencia laboral (NÚMERO ENTERO, redondear al más cercano)")
    experiencia_en_empresa_actual: int = Field(0, description="Años en la empresa actual o más reciente (NÚMERO ENTERO, redondear al más cercano)")
    exp_cargo_actual: int = Field(0, description="Años en el cargo actual o más reciente (NÚMERO ENTERO, redondear al más cercano)")
    exp_proy_similares: int = Field(0, description="Años de experiencia en proyectos industriales/mineros similares al cargo objetivo (NÚMERO ENTERO, redondear al más cercano)")

    antecedentes_academicos: str = Field("", description="Antecedentes académicos. Formato: 'TITULO - INSTITUCION - ESTADO (Titulado/Egresado/En Curso)', uno por línea.")
    experiencia_general: str = Field("", description="Experiencia General para informe TEC-02. NUNCA resumir. Reconstruir TODA la línea de tiempo. Formato EXACTO por entrada: 'YYYY-YYYY CARGO - EMPRESA - FAENA/ESTABLECIMIENTO'. Una entrada por línea.")
    experiencia_especifica: str = Field("", description="Experiencia Específica para informe TEC-02A. Enfocarse en roles de PROTECCIÓN CONTRA INCENDIOS (SPCI), MINERÍA o proyectos INDUSTRIALES. Mismo formato: 'YYYY-YYYY CARGO - EMPRESA - FAENA'.")
    otras_experiencias: str = Field("", description="Otras experiencias relevantes, certificaciones, cursos. Mismo formato de lista.")

    evaluacion_general: str = Field("", description="Resumen táctico profesional en 3 párrafos: 1. Propuesta de Valor, 2. Dominio Técnico, 3. Aptitud Operacional. En español formal de Chile.")
    match_score: int = Field(80, description="Índice de mérito de 0 a 100 basado en antigüedad, certificaciones y habilidades técnicas especializadas.")
    match_explicacion: str = Field("", description="Justificación clara y profesional del match_score. Destacar brechas o fortalezas técnicas.")
    nota: float = Field(1.0, description="Nota numérica de 1.0 a 7.0 (escala chilena). Basada en: 30% antigüedad, 40% certificaciones, 30% complejidad de proyectos.")
    ranking: int = Field(50, description="Ranking estratégico de 1 a 100 representando la competitividad del candidato en el sector industrial.")


# --- Core Processor ---

class AFKProcessor:
    def __init__(self):
        self.supabase_url = os.getenv("SUPABASE_URL")
        self.supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
        self.openai_api_key = os.getenv("OPENAI_API_KEY")

        if not all([self.supabase_url, self.supabase_key, self.openai_api_key]):
            print("WARNING: Missing environment variables. Ensure SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, and OPENAI_API_KEY are set.")

        self.supabase: Client = create_client(self.supabase_url, self.supabase_key) if self.supabase_url else None
        self.openai = OpenAI(api_key=self.openai_api_key) if self.openai_api_key else None

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extracción de texto con Layout-Aware usando pdfplumber (preferido).
        Fallback a pypdf si pdfplumber no está disponible.
        """
        try:
            import pdfplumber
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    # Extraer palabras con posición (x, y, texto)
                    words = page.extract_words(
                        x_tolerance=5,
                        y_tolerance=5,
                        keep_blank_chars=False,
                        use_text_flow=False,
                        extra_attrs=["size", "fontname"]
                    )

                    if not words:
                        # Fallback a extracción simple si no hay palabras con layout
                        simple_text = page.extract_text() or ""
                        full_text += simple_text + f"\n[PÁGINA_BREAK_{i+1}]\n"
                        continue

                    # Agrupar palabras en líneas por proximidad en Y
                    lines = []
                    current_line = []
                    last_y = None
                    Y_THRESHOLD = 3  # tolerancia en puntos para agrupar en la misma línea

                    # Ordenar por Y descendente (top), luego X ascendente (left)
                    sorted_words = sorted(words, key=lambda w: (-round(w["top"] / Y_THRESHOLD), w["x0"]))

                    for word in sorted_words:
                        y_bucket = round(word["top"] / Y_THRESHOLD)
                        if last_y is None or abs(word["top"] - last_y) <= Y_THRESHOLD * 2:
                            current_line.append(word["text"])
                            last_y = word["top"]
                        else:
                            if current_line:
                                lines.append(" ".join(current_line))
                            current_line = [word["text"]]
                            last_y = word["top"]

                    if current_line:
                        lines.append(" ".join(current_line))

                    page_text = "\n".join(lines)
                    full_text += page_text + f"\n[PÁGINA_BREAK_{i+1}]\n"

            return full_text

        except ImportError:
            # Fallback a pypdf si pdfplumber no está instalado
            print("⚠️  pdfplumber no disponible. Usando extracción básica con pypdf.")
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except Exception as e:
            print(f"⚠️  Error con pdfplumber ({e}). Intentando con pypdf...")
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text

    def extract_text_from_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Incluir tablas también
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)

    def extract_text(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower().strip()
        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Formato de archivo no soportado: '{ext}'")

    def _pdf_to_base64_images(self, file_path: str, max_pages: int = 3) -> List[str]:
        """Convierte las primeras páginas de un PDF a imágenes base64 usando PyMuPDF."""
        base64_images = []
        try:
            doc = fitz.open(file_path)
            for page_num in range(min(len(doc), max_pages)):
                page = doc.load_page(page_num)
                # Renderizar página a imagen (zoom 2x para mejor resolución)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("jpeg")
                b64_str = base64.b64encode(img_bytes).decode("utf-8")
                base64_images.append(b64_str)
            doc.close()
        except Exception as e:
            print(f"❌ Error convirtiendo PDF a imágenes: {e}")
        return base64_images

    def process_cv_with_ai(self, text: str, filename: str = "", file_path: str = None) -> CandidateCV:
        use_vision = False
        base64_images = []

        if len(text.strip()) < 50:
            if file_path and file_path.lower().endswith(".pdf"):
                print("⚠️ Texto insuficiente detectado. Activando Plan B: OCR con GPT-4o Vision API...")
                base64_images = self._pdf_to_base64_images(file_path)
                if not base64_images:
                    raise ValueError("Texto extraído insuficiente y falló la conversión a imágenes. Abortando.")
                use_vision = True
            else:
                raise ValueError("Texto extraído insuficiente y no es un PDF válido para OCR visual. Abortando para evitar alucinaciones de IA.")

        system_prompt = (
            "Eres JARVIS (Just A Rather Very Intelligent System), Analista Táctico de RRHH de élite de Stark Industries. "
            "REGLA CRÍTICA DE IDIOMA: TODOS los campos de salida DEBEN estar en ESPAÑOL (Español Profesional de Chile). "
            "REGLA CRÍTICA DE IDENTIDAD: NUNCA inventes nombres genéricos como 'Juan Pérez'. Si no encuentras el nombre en el texto, usa el nombre del 'Archivo original'. "
            "NUNCA uses inglés para resúmenes, cargos o evaluaciones. "
            "Tu objetivo es realizar una extracción EXHAUSTIVA, PROFESIONAL y SIN PÉRDIDA DE DATOS del CV proporcionado. "
            "TONO: Formal, táctico y eficiente. Usa terminología de alta precisión. "
            "PROTOCOLO STARK ABSOLUTO v4 (Cero pérdida de datos): "
            "1. HISTORIAL PROFESIONAL: NUNCA resumir. DEBES reconstruir la línea de tiempo COMPLETA entrada por entrada. "
            "2. FORMATO OBLIGATORIO: Para 'experiencia_general' y 'experiencia_especifica', usa EXACTAMENTE: 'YYYY-YYYY CARGO - EMPRESA - FAENA/ESTABLECIMIENTO' "
            "(Ejemplo: '2017-2019 OPERADOR TRACTO CAMIÓN - GREKAT - COLLAHUASI'). Una entrada por línea. "
            "3. IDENTIFICACIÓN DE FAENAS: Auditar el texto en busca de faenas mineras o ubicaciones de proyecto e incluirlas en la entrada. "
            "4. REGISTROS ACADÉMICOS: Formato: 'TITULO - INSTITUCION - ESTADO (Titulado, Egresado, o En Curso)'. "
            "5. INVENTARIO DE SOFTWARE: Lista detallada de ERPs (SAP, Maximo), herramientas técnicas y software específico. "
            "6. EVALUACION GENERAL: Resumen de alta fidelidad en 3 párrafos: Propuesta de Valor, Dominio Técnico, Aptitud Operacional. "
            "7. CALIFICACIÓN: 'nota' (1.0-7.0) y 'ranking' (1-100) deben reflejar el rigor del sector industrial/minero.\n"
            "8. NÚMEROS ENTEROS: Todos los campos de 'experiencia_total', 'experiencia_en_empresa_actual', 'exp_cargo_actual' y 'exp_proy_similares' DEBEN ser NÚMEROS ENTEROS. Redondea fracciones al entero más cercano (ej: 2.92 -> 3).\n"
            f"9. AÑO ACTUAL: Asume que hoy es el año {datetime.now().year}. Calcula los tiempos de experiencia 'hasta la fecha' (Presente) usando este año."
        )

        tools = [
            {
                "type": "function",
                "function": {
                    "name": "extract_cv_data",
                    "description": "Extrae información estructurada de un CV de forma exhaustiva siguiendo el Protocolo Stark Absoluto v4",
                    "parameters": CandidateCV.model_json_schema()
                }
            }
        ]

        messages = [
            {"role": "system", "content": system_prompt}
        ]

        if use_vision:
            # Construir mensaje con imágenes para GPT-4o Vision
            user_content = [
                {"type": "text", "text": f"Archivo original: {filename}\n\nEste CV no contiene texto seleccionable. Aquí están las imágenes de las páginas para que las leas:"}
            ]
            for b64_img in base64_images:
                user_content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{b64_img}",
                        "detail": "high"
                    }
                })
            messages.append({"role": "user", "content": user_content})
        else:
            messages.append({"role": "user", "content": f"Archivo original: {filename}\n\nTexto extraído:\n{text[:55000]}"})
        
        response = self.openai.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            tools=tools,
            tool_choice={"type": "function", "function": {"name": "extract_cv_data"}}
        )

        tool_call = response.choices[0].message.tool_calls[0]
        data = json.loads(tool_call.function.arguments)

        return CandidateCV(**data)

    def generate_embedding(self, text: str) -> List[float]:
        # Limitar texto para embedding (máx 8191 tokens)
        # Reducido de 30000 a 15000 caracteres para asegurar que no exceda los 8192 tokens
        truncated = text[:15000]
        try:
            response = self.openai.embeddings.create(
                input=truncated,
                model="text-embedding-3-small"
            )
            return response.data[0].embedding
        except Exception as e:
            print(f"⚠️ Error al generar embedding, reduciendo texto: {e}")
            truncated_fallback = text[:8000]
            response_fallback = self.openai.embeddings.create(
                input=truncated_fallback,
                model="text-embedding-3-small"
            )
            return response_fallback.data[0].embedding

    def sync_to_supabase(self, candidate_id: Optional[str], cv_data: CandidateCV, full_text: str):
        if not self.supabase:
            print("Supabase client not initialized. Skipping sync.")
            return

        # 1. Generar Embedding
        print("🧠 Generando embedding vectorial...")
        embedding = self.generate_embedding(full_text)

        # 2. Construir payload con columnas que existen en la tabla candidates
        EXISTING_DB_COLUMNS = {
            "nombre_completo", "rut", "fecha_nacimiento", "correo", "telefono",
            "profesion", "experiencia", "software_que_domina", "ultima_exp_laboral_empresa",
            "cargo", "periodo", "experiencia_general", "experiencia_especifica",
            "otras_experiencias", "evaluacion_general", "experiencia_total",
            "experiencia_en_empresa_actual", "exp_cargo_actual", "exp_proy_similares",
            "cargo_a_desempenar", "nota", "ranking", "status", "vacancy_id",
            "match_score", "onboarding_progress", "source",
            "antecedentes_academicos", "direccion", "cv_full_text", "cv_embedding", "resumen_ia"
        }

        raw = cv_data.model_dump()
        payload = {k: v for k, v in raw.items() if k in EXISTING_DB_COLUMNS}
        payload["status"] = "Analizado por IA"
        payload["cv_full_text"] = full_text
        payload["cv_embedding"] = embedding

        # Sanitizar strings (PostgreSQL no soporta el carácter nulo \u0000)
        for key, val in payload.items():
            if isinstance(val, str):
                payload[key] = val.replace('\x00', '')

        # match_score es integer en Supabase
        if "match_score" in payload and payload["match_score"] is not None:
            payload["match_score"] = int(payload["match_score"])

        if candidate_id:
            print(f"🔄 Actualizando candidato ID: {candidate_id}...")
            result = self.supabase.table("candidates").update(payload).eq("id", candidate_id).execute()
        else:
            print("🆕 Sin ID. Verificando candidato existente por RUT o Nombre...")
            found_id = None

            # Verificar por RUT primero
            if cv_data.rut:
                clean_rut = cv_data.rut.replace(" ", "").strip()
                existing = self.supabase.table("candidates").select("id").eq("rut", clean_rut).execute()
                if existing.data:
                    found_id = existing.data[0]['id']

            # Verificar por Nombre si no se encontró por RUT
            if not found_id and cv_data.nombre_completo:
                existing = self.supabase.table("candidates").select("id").ilike("nombre_completo", cv_data.nombre_completo.strip()).execute()
                if existing.data:
                    found_id = existing.data[0]['id']

            if found_id:
                print(f"⚠️ Candidato {cv_data.nombre_completo} ya existe (ID: {found_id}). Actualizando.")
                result = self.supabase.table("candidates").update(payload).eq("id", found_id).execute()
            else:
                print(f"✨ Insertando nuevo candidato: {cv_data.nombre_completo}")
                result = self.supabase.table("candidates").insert(payload).execute()

        print(f"✅ Sync completado para: {cv_data.nombre_completo}")
        return result


# --- CLI for Testing ---

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="AFK Intelligent CV Processor — Protocolo Stark v4")
    parser.add_argument("file", help="Ruta al archivo CV (PDF o DOCX)")
    parser.add_argument("--id", help="ID del candidato en Supabase (opcional: actualiza si existe)", required=False)
    args = parser.parse_args()

    processor = AFKProcessor()

    print(f"🚀 Iniciando procesamiento: {args.file}")
    try:
        print("📄 Extrayendo texto con motor layout-aware...")
        text = processor.extract_text(args.file)
        char_count = len(text)
        print(f"📝 Texto extraído: {char_count:,} caracteres. Enviando a JARVIS...")

        cv_data = processor.process_cv_with_ai(text, filename=os.path.basename(args.file), file_path=args.file)
        print(f"✅ Extracción IA completada: {cv_data.nombre_completo} | Nota: {cv_data.nota} | Ranking: {cv_data.ranking}")

        print("🔄 Sincronizando con Supabase...")
        processor.sync_to_supabase(args.id, cv_data, text)
        print("🎯 Protocolo Stark: OPERACIÓN COMPLETADA.")

    except Exception as e:
        import traceback
        print(f"❌ Error durante el procesamiento: {e}")
        traceback.print_exc()
